# Увоз потребних библиотека за рад са Word и Excel датотекама
# python-docx за читање Word докумената
# pandas за манипулацију подацима и креирање Excel датотека
# openpyxl за напредно форматирање Excel ћелија
try:
    import docx
    import pandas as pd
    from pathlib import Path
    from openpyxl.styles import Alignment, PatternFill, Font
except ImportError:
    print("Потребни пакети нису инсталирани. Молимо покрените:")
    print("pip install -r requirements.txt")
    exit(1)

# Функција која израчунава оптималну ширину колоне на основу садржаја
# Обрада различитих типова података за израчунавање ширине колоне
def get_column_width(col_series):
    max_length = 0
    for value in col_series:
        try:
            # Конверзија вредности у текст и добијање дужине
            str_value = str(value) if value is not None else ''
            max_length = max(max_length, len(str_value))
        except:
            continue
    return max_length

# Функција која проверава да ли је текст подебљан (bold)
# Проверава различите начине означавања подебљаног текста у Word документу
def is_bold_run(run):
    try:
        return run.bold or (hasattr(run, 'font') and run.font.name and 'bold' in run.font.name.lower())
    except AttributeError:
        return False

# Функција за конверзију текстуалних бројева у нумерички формат
# Подржава европски формат бројева (нпр. 200.000,00) и стандардни формат
def try_convert_number(text):
    """Покушај конверзије текста у број, обрада различитих формата"""
    # Уклањање размака али задржавање тачака и зареза
    cleaned = text.strip()
    
    try:
        if '.' in cleaned and ',' in cleaned:
            # Европски формат (нпр. 200.000,00)
            cleaned = cleaned.replace('.', '')  # Уклањање сепаратора хиљада
            cleaned = cleaned.replace(',', '.')  # Конверзија децималног сепаратора
        
        # Конверзија у float и заокруживање на 2 децимале
        value = round(float(cleaned.replace(' ', '')), 2)
        return value
    except (ValueError, TypeError):
        return text

# Функција за форматирање нумеричких вредности у Excel ћелијама
# Примена форматирања бројева на ћелију
def format_number(cell, value):
    """Примена форматирања бројева на ћелију"""
    if isinstance(value, (int, float)):
        # Коришћење стандардног формата са зарезом као сепаратором хиљада и тачком за децимале
        cell.number_format = '#,##0.00'  # Приказаће: 200,000.00
        cell.alignment = Alignment(
            vertical='top',
            horizontal='center',  # Центрирано поравнање
            wrap_text=False,
            shrink_to_fit=False
        )
        return True
    return False

# Функција која извлачи текст и информације о форматирању из Word елемената
# Обрада параграфа и ћелија табеле за подебљани текст
def get_text_with_format(text_element):
    """Обрада параграфа и ћелија табеле за подебљани текст"""
    text = ''
    is_bold = False
    
    # Обрада једног параграфа наспрам више параграфа (ћелије табеле)
    paragraphs = text_element.paragraphs if hasattr(text_element, 'paragraphs') else [text_element]
    
    for paragraph in paragraphs:
        para_text = ''
        para_bold = False
        
        # Прво провера директног стила параграфа
        try:
            if paragraph.style and paragraph.style.font and paragraph.style.font.bold:
                para_bold = True
        except AttributeError:
            pass
        
        # Провера појединачних секција
        try:
            if hasattr(paragraph, 'runs') and paragraph.runs:
                for run in paragraph.runs:
                    run_text = run.text.strip()
                    if run_text:
                        # Провера директног подебљања и својстава фонта
                        if getattr(run, 'bold', False) or (hasattr(run, 'font') and getattr(run.font, 'bold', False)):
                            para_bold = True
                        para_text += run.text
            else:
                para_text = paragraph.text
        except AttributeError:
            para_text = paragraph.text if hasattr(paragraph, 'text') else ''
        
        text += para_text
        if para_bold:
            is_bold = True
    
    return text.strip(), is_bold

# Функција која форматира текст тако да број буде у новом реду после имена
# Форматирање текста тако да број буде у новом реду после имена
def format_name_number(text):
    """Форматирање текста тако да број буде у новом реду после имена"""
    if not text:
        return text
        
    # Проналажење позиције последњег словног карактера
    last_letter_pos = -1
    for i, char in enumerate(text):
        if char.isalpha():
            last_letter_pos = i
            
    # Ако смо пронашли слово и није на крају
    if last_letter_pos >= 0 and last_letter_pos < len(text) - 1:
        name_part = text[:last_letter_pos + 1].strip()
        number_part = text[last_letter_pos + 1:].strip()
        if name_part and number_part:
            return f"{name_part}\n{number_part}"
            
    return text

# Главна функција за конверзију Word документа у Excel
# Обрађује текст и табеле, задржава форматирање и структуру документа
def word_to_excel(word_file, excel_file):
    # Учитавање Word документа у меморију
    doc = docx.Document(word_file)
    
    # Иницијализација листи за чување података и информација о форматирању
    rows = []
    is_bold_format = []  # Праћење подебљаног текста
    
    row_count = 0  # Бројач редова у табели
    
    # Секвенцијална обрада документа
    for element in doc.element.body:
        if element.tag.endswith('tbl'):
            table = doc.tables[len([e for e in doc.element.body[:doc.element.body.index(element)] 
                                  if e.tag.endswith('tbl')])]
            
            rows.append([''])
            is_bold_format.append([False])
            
            for row in table.rows:
                row_data = []
                row_format = []
                row_count += 1
                
                for col_idx, cell in enumerate(row.cells):
                    text, is_bold = get_text_with_format(cell)
                    # Промена текста четврте колоне ако је први ред
                    if col_idx == 3 and row_count == 1:
                        text = "Текући рачун"
                    elif col_idx == 1:
                        text = format_name_number(text)
                    elif col_idx == 2:
                        text = try_convert_number(text.strip())
                    row_data.append(text)
                    row_format.append(is_bold)
                
                if any(str(x) for x in row_data):
                    rows.append(row_data)
                    is_bold_format.append(row_format)
            
            # Додавање празног реда после табеле
            rows.append([''])
            is_bold_format.append([False])
                
        elif element.tag.endswith('p'):  # Параграф
            text, is_bold = get_text_with_format(element)
            if text:
                rows.append([text])
                is_bold_format.append([is_bold])
    
    # Креирање DataFrame-а без заглавља
    final_df = pd.DataFrame(rows)
    
    # Додавање колоне маркера за идентификацију редова табеле
    is_table = []
    for row in rows:
        is_table.append(len(row) > 1)  # Ако ред има више колона, то је ред табеле
    
    # Чување у Excel са форматирањем
    try:
        with pd.ExcelWriter(excel_file, engine='openpyxl') as writer:
            # Упис DataFrame-а без заглавља
            final_df.to_excel(writer, 
                            sheet_name='Document Content',
                            index=False,
                            header=False)
            worksheet = writer.sheets['Document Content']
            
            # Подешавање ширине колона
            for idx in range(final_df.shape[1]):
                col_letter = chr(65 + idx) if idx < 26 else chr(64 + (idx//26)) + chr(65 + (idx%26))
                if idx in [1, 2, 3]:  # Колоне 2, 3 и 4
                    worksheet.column_dimensions[col_letter].width = 30  # Шире колоне
                else:
                    worksheet.column_dimensions[col_letter].width = 20  # Подразумевана ширина
            
            # Подешавање особина ћелија са Calibri фонтом
            for row_idx, row in enumerate(worksheet.iter_rows()):
                for col_idx, cell in enumerate(row):
                    # Основно поравнање за све ћелије
                    align_props = {
                        'vertical': 'top',
                        'horizontal': 'left',
                        'wrap_text': True if col_idx == 1 else False,  # Омогућавање преламања текста за другу колону
                        'shrink_to_fit': False
                    }
                    
                    # Обрада форматирања бројева у трећој колони
                    if col_idx == 2:
                        if not format_number(cell, cell.value):
                            cell.alignment = Alignment(**align_props)
                    else:
                        cell.alignment = Alignment(**align_props)
                    
                    try:
                        is_bold = is_bold_format[row_idx][col_idx]
                    except IndexError:
                        is_bold = False
                        
                    cell.font = Font(
                        name='Calibri',
                        size=11 if is_table[row_idx] else 12,
                        bold=is_bold
                    )
                
                # Increase height for rows with wrapped text
                if any(cell.alignment.wrap_text for cell in row):
                    worksheet.row_dimensions[row[0].row].height = 30
                else:
                    worksheet.row_dimensions[row[0].row].height = 15
    
    except Exception as e:
        print(f"Упозорење: Грешка при чувању датотеке: {str(e)}")
        return False
    
    return True

# Главни део програма
# Проналази све Word документе у тренутном директоријуму и конвертује их у Excel
if __name__ == "__main__":
    # Добијање тренутног директоријума
    current_dir = Path('.')
    # Креирање излазног директоријума ако не постоји
    output_dir = current_dir / 'ex'
    output_dir.mkdir(exist_ok=True)
    
    # Проналажење свих .docx датотека
    docx_files = list(current_dir.glob('*.docx'))
    
    if not docx_files:
        print("Нема .docx датотека у тренутном директоријуму!")
        exit(1)
        
    print(f"Пронаћено {len(docx_files)} Word датотека за обраду...")
    
    for word_path in docx_files:
        try:
            # Креирање излазне путање са истим именом али .xlsx екстензијом у 'ex' фолдеру
            excel_path = output_dir / f"{word_path.stem}.xlsx"
            
            success = word_to_excel(word_path, excel_path)
            if success:
                print(f"Конвертовано: {word_path.name} → {excel_path.name}")
        except Exception as e:
            print(f"Грешка при конвертовању {word_path.name}: {str(e)}")
    
    print("\nОбрада завршена! Проверите директоријум 'ex' за излазне датотеке.")
