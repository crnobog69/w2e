try:
    import docx
    import pandas as pd
    from pathlib import Path
    from openpyxl.styles import Alignment, PatternFill, Font
except ImportError:
    print("Потребни пакети нису инсталирани. Молимо покрените:")
    print("pip install -r requirements.txt")
    exit(1)

def get_column_width(col_series):
    # Handle different data types for column width calculation
    max_length = 0
    for value in col_series:
        try:
            # Convert any value to string and get its length
            str_value = str(value) if value is not None else ''
            max_length = max(max_length, len(str_value))
        except:
            continue
    return max_length

def is_bold_run(run):
    try:
        return run.bold or (hasattr(run, 'font') and run.font.name and 'bold' in run.font.name.lower())
    except AttributeError:
        return False

def try_convert_number(text):
    """Try to convert text to number, handling various formats"""
    # Remove any whitespace but keep dots and commas
    cleaned = text.strip()
    
    try:
        if '.' in cleaned and ',' in cleaned:
            # European format (e.g. 200.000,00)
            cleaned = cleaned.replace('.', '')  # Remove thousand separators
            cleaned = cleaned.replace(',', '.')  # Convert decimal separator
        
        # Convert to float and round to 2 decimal places
        value = round(float(cleaned.replace(' ', '')), 2)
        return value
    except (ValueError, TypeError):
        return text

def format_number(cell, value):
    """Apply number formatting to cell"""
    if isinstance(value, (int, float)):
        # Use standard format with comma as thousand separator and dot for decimals
        cell.number_format = '#,##0.00'  # This will show: 200,000.00
        cell.alignment = Alignment(
            vertical='top',
            horizontal='center',  # Changed to center
            wrap_text=False,
            shrink_to_fit=False
        )
        return True
    return False

def get_text_with_format(text_element):
    """Handle both paragraphs and table cells for bold text"""
    text = ''
    is_bold = False
    
    # Handle single paragraph vs multiple paragraphs (table cells)
    paragraphs = text_element.paragraphs if hasattr(text_element, 'paragraphs') else [text_element]
    
    for paragraph in paragraphs:
        para_text = ''
        para_bold = False
        
        # Check direct paragraph style first
        try:
            if paragraph.style and paragraph.style.font and paragraph.style.font.bold:
                para_bold = True
        except AttributeError:
            pass
        
        # Check individual runs
        try:
            if hasattr(paragraph, 'runs') and paragraph.runs:
                for run in paragraph.runs:
                    run_text = run.text.strip()
                    if run_text:
                        # Check both direct bold and font properties
                        if getattr(run, 'bold', False) or (hasattr(run, 'font') and getattr(run.font, 'bold', False)):
                            para_bold = True
                        para_text += run.text
            else:
                para_text = paragraph.text
        except AttributeError:
            para_text = paragraph.text if hasattr(paragraph, 'text') else ''
        
        text += para_text
        if para_bold:
            is_bold = True
    
    return text.strip(), is_bold

def format_name_number(text):
    """Format text to put number on new line after name by finding last letter"""
    if not text:
        return text
        
    # Find the last alphabetic character position
    last_letter_pos = -1
    for i, char in enumerate(text):
        if char.isalpha():
            last_letter_pos = i
            
    # If we found a letter and it's not at the end
    if last_letter_pos >= 0 and last_letter_pos < len(text) - 1:
        name_part = text[:last_letter_pos + 1].strip()
        number_part = text[last_letter_pos + 1:].strip()
        if name_part and number_part:
            return f"{name_part}\n{number_part}"
            
    return text

def word_to_excel(word_file, excel_file):
    # Load the Word document
    doc = docx.Document(word_file)
    
    # Initialize data storage with single column
    rows = []
    is_bold_format = []  # Track bold formatting
    
    row_count = 0  # Track row number within table
    
    # Process document sequentially
    for element in doc.element.body:
        if element.tag.endswith('tbl'):
            table = doc.tables[len([e for e in doc.element.body[:doc.element.body.index(element)] 
                                  if e.tag.endswith('tbl')])]
            
            rows.append([''])
            is_bold_format.append([False])
            
            for row in table.rows:
                row_data = []
                row_format = []
                row_count += 1
                
                for col_idx, cell in enumerate(row.cells):
                    text, is_bold = get_text_with_format(cell)
                    # Change fourth column text if it's the first row
                    if col_idx == 3 and row_count == 1:
                        text = "Текући рачун"
                    elif col_idx == 1:
                        text = format_name_number(text)
                    elif col_idx == 2:
                        text = try_convert_number(text.strip())
                    row_data.append(text)
                    row_format.append(is_bold)
                
                if any(str(x) for x in row_data):
                    rows.append(row_data)
                    is_bold_format.append(row_format)
            
            # Add blank line after table
            rows.append([''])
            is_bold_format.append([False])
                
        elif element.tag.endswith('p'):  # Paragraph
            text, is_bold = get_text_with_format(element)
            if text:
                rows.append([text])
                is_bold_format.append([is_bold])
    
    # Create DataFrame without headers
    final_df = pd.DataFrame(rows)
    
    # Add a marker column to identify table rows
    is_table = []
    for row in rows:
        is_table.append(len(row) > 1)  # If row has multiple columns, it's a table
    
    # Save to Excel with formatting
    try:
        with pd.ExcelWriter(excel_file, engine='openpyxl') as writer:
            # Write DataFrame without headers
            final_df.to_excel(writer, 
                            sheet_name='Document Content',
                            index=False,
                            header=False)
            worksheet = writer.sheets['Document Content']
            
            # Set column widths
            for idx in range(final_df.shape[1]):
                col_letter = chr(65 + idx) if idx < 26 else chr(64 + (idx//26)) + chr(65 + (idx%26))
                if idx in [1, 2, 3]:  # Columns 2, 3, and 4
                    worksheet.column_dimensions[col_letter].width = 30  # Wider columns
                else:
                    worksheet.column_dimensions[col_letter].width = 20  # Default width
            
            # Set cell properties with Calibri font
            for row_idx, row in enumerate(worksheet.iter_rows()):
                for col_idx, cell in enumerate(row):
                    # Basic alignment for all cells
                    align_props = {
                        'vertical': 'top',
                        'horizontal': 'left',
                        'wrap_text': True if col_idx == 1 else False,  # Enable wrap for second column
                        'shrink_to_fit': False
                    }
                    
                    # Handle number formatting in the third column
                    if col_idx == 2:
                        if not format_number(cell, cell.value):
                            cell.alignment = Alignment(**align_props)
                    else:
                        cell.alignment = Alignment(**align_props)
                    
                    try:
                        is_bold = is_bold_format[row_idx][col_idx]
                    except IndexError:
                        is_bold = False
                        
                    cell.font = Font(
                        name='Calibri',
                        size=11 if is_table[row_idx] else 12,
                        bold=is_bold
                    )
                
                # Increase height for rows with wrapped text
                if any(cell.alignment.wrap_text for cell in row):
                    worksheet.row_dimensions[row[0].row].height = 30
                else:
                    worksheet.row_dimensions[row[0].row].height = 15
    
    except Exception as e:
        print(f"Упозорење: Грешка при чувању датотеке: {str(e)}")
        return False
    
    return True

if __name__ == "__main__":
    # Get current directory
    current_dir = Path('.')
    # Create output directory if it doesn't exist
    output_dir = current_dir / 'ex'
    output_dir.mkdir(exist_ok=True)
    
    # Find all .docx files
    docx_files = list(current_dir.glob('*.docx'))
    
    if not docx_files:
        print("Нема .docx датотека у тренутном директоријуму!")
        exit(1)
        
    print(f"Пронаћено {len(docx_files)} Word датотека за обраду...")
    
    for word_path in docx_files:
        try:
            # Create output path with same name but .xlsx extension in 'ex' folder
            excel_path = output_dir / f"{word_path.stem}.xlsx"
            
            success = word_to_excel(word_path, excel_path)
            if success:
                print(f"Конвертовано: {word_path.name} → {excel_path.name}")
        except Exception as e:
            print(f"Грешка при конвертовању {word_path.name}: {str(e)}")
    
    print("\nОбрада завршена! Проверите директоријум 'ex' за излазне датотеке.")
