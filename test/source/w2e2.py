try:
    import docx
    import pandas as pd
    from pathlib import Path
    from openpyxl.styles import Alignment, PatternFill, Font
except ImportError:
    print("Required packages are not installed. Please run:")
    print("pip install -r requirements.txt")
    exit(1)

def get_column_width(col_series):
    # Handle different data types for column width calculation
    max_length = 0
    for value in col_series:
        try:
            # Convert any value to string and get its length
            str_value = str(value) if value is not None else ''
            max_length = max(max_length, len(str_value))
        except:
            continue
    return max_length

def is_bold_run(run):
    try:
        return run.bold or (hasattr(run, 'font') and run.font.name and 'bold' in run.font.name.lower())
    except AttributeError:
        return False

def try_convert_number(text):
    """Try to convert text to number, handling various formats"""
    # Remove any whitespace and handle European number format
    cleaned = text.strip().replace(' ', '')
    if ',' in cleaned and '.' not in cleaned:
        # Handle European format (e.g., "200.000,00" -> 200000.00)
        cleaned = cleaned.replace('.', '').replace(',', '.')
    else:
        # Handle regular format with comma thousands separator
        cleaned = cleaned.replace(',', '')
    
    try:
        if '.' in cleaned:
            return float(cleaned)
        return int(cleaned)
    except (ValueError, TypeError):
        return text

def format_number(cell, value):
    """Apply number formatting to cell"""
    if isinstance(value, (int, float)):
        cell.number_format = '#,##0.00'
        cell.alignment = Alignment(
            vertical='top',
            horizontal='right',
            wrap_text=False,
            shrink_to_fit=False
        )
        return True
    return False

def get_text_with_format(text_element):
    """Handle both paragraphs and table cells for bold text"""
    text = ''
    is_bold = False
    
    # Handle single paragraph vs multiple paragraphs (table cells)
    paragraphs = text_element.paragraphs if hasattr(text_element, 'paragraphs') else [text_element]
    
    for paragraph in paragraphs:
        para_text = ''
        para_bold = False
        
        # Check direct paragraph style first
        try:
            if paragraph.style and paragraph.style.font and paragraph.style.font.bold:
                para_bold = True
        except AttributeError:
            pass
        
        # Check individual runs
        try:
            if hasattr(paragraph, 'runs') and paragraph.runs:
                for run in paragraph.runs:
                    run_text = run.text.strip()
                    if run_text:
                        # Check both direct bold and font properties
                        if getattr(run, 'bold', False) or (hasattr(run, 'font') and getattr(run.font, 'bold', False)):
                            para_bold = True
                        para_text += run.text
            else:
                para_text = paragraph.text
        except AttributeError:
            para_text = paragraph.text if hasattr(paragraph, 'text') else ''
        
        text += para_text
        if para_bold:
            is_bold = True
    
    return text.strip(), is_bold

def word_to_excel(word_file, excel_file):
    # Load the Word document
    doc = docx.Document(word_file)
    
    # Initialize data storage with single column
    rows = []
    is_bold_format = []  # Track bold formatting
    
    # Process document sequentially
    for element in doc.element.body:
        if element.tag.endswith('tbl'):  # Table
            table = doc.tables[len([e for e in doc.element.body[:doc.element.body.index(element)] 
                                  if e.tag.endswith('tbl')])]
            
            # Add blank line before table
            rows.append([''])
            is_bold_format.append([False])
            
            for row in table.rows:
                row_data = []
                row_format = []
                for col_idx, cell in enumerate(row.cells):
                    text, is_bold = get_text_with_format(cell)
                    # Convert numbers in the third column
                    if col_idx == 2:
                        text = try_convert_number(text.strip())
                    row_data.append(text)
                    row_format.append(is_bold)
                if any(str(x) for x in row_data):  # Check non-empty rows
                    rows.append(row_data)
                    is_bold_format.append(row_format)
            
            # Add blank line after table
            rows.append([''])
            is_bold_format.append([False])
                
        elif element.tag.endswith('p'):  # Paragraph
            text, is_bold = get_text_with_format(element)
            if text:
                rows.append([text])
                is_bold_format.append([is_bold])
    
    # Create DataFrame without headers
    final_df = pd.DataFrame(rows)
    
    # Add a marker column to identify table rows
    is_table = []
    for row in rows:
        is_table.append(len(row) > 1)  # If row has multiple columns, it's a table
    
    # Save to Excel with formatting
    try:
        with pd.ExcelWriter(excel_file, engine='openpyxl') as writer:
            # Write DataFrame without headers
            final_df.to_excel(writer, 
                            sheet_name='Document Content',
                            index=False,
                            header=False)
            worksheet = writer.sheets['Document Content']
            
            # Set default column width
            for idx in range(final_df.shape[1]):
                col_letter = chr(65 + idx) if idx < 26 else chr(64 + (idx//26)) + chr(65 + (idx%26))
                worksheet.column_dimensions[col_letter].width = 20  # Fixed width
            
            # Set cell properties with Calibri font
            for row_idx, row in enumerate(worksheet.iter_rows()):
                for col_idx, cell in enumerate(row):
                    # Basic alignment for all cells
                    align_props = {
                        'vertical': 'top',
                        'horizontal': 'left',
                        'wrap_text': False,
                        'shrink_to_fit': False
                    }
                    
                    # Handle number formatting in the third column
                    if col_idx == 2:
                        if not format_number(cell, cell.value):
                            cell.alignment = Alignment(**align_props)
                    else:
                        cell.alignment = Alignment(**align_props)
                    
                    try:
                        is_bold = is_bold_format[row_idx][col_idx]
                    except IndexError:
                        is_bold = False
                        
                    cell.font = Font(
                        name='Calibri',
                        size=11 if is_table[row_idx] else 12,
                        bold=is_bold
                    )
                
                worksheet.row_dimensions[row[0].row].height = 15
    
    except Exception as e:
        print(f"Warning: Error while saving file: {str(e)}")
        return False
    
    return True

if __name__ == "__main__":
    # Get current directory
    current_dir = Path('.')
    # Create output directory if it doesn't exist
    output_dir = current_dir / 'ex'
    output_dir.mkdir(exist_ok=True)
    
    # Find all .docx files
    docx_files = list(current_dir.glob('*.docx'))
    
    if not docx_files:
        print("No .docx files found in the current directory!")
        exit(1)
        
    print(f"Found {len(docx_files)} Word files to process...")
    
    for word_path in docx_files:
        try:
            # Create output path with same name but .xlsx extension in 'ex' folder
            excel_path = output_dir / f"{word_path.stem}.xlsx"
            
            success = word_to_excel(word_path, excel_path)
            if success:
                print(f"Converted: {word_path.name} → {excel_path.name}")
        except Exception as e:
            print(f"Error converting {word_path.name}: {str(e)}")
    
    print("\nProcessing complete! Check the 'ex' folder for output files.")
